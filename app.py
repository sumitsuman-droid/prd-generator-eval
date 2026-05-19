import streamlit as st
import re
import json
from io import BytesIO

from docx import Document
from generator import generate_prd_streaming, evaluate_prd, parse_prd_text, explain_retrieval
from retrieval import get_similar_prd_metas, retrieve_similar_prds


def _slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")


def _streaming_display_text(raw_json_partial: str) -> str:
    """Extract completed top-level JSON fields and render as clean markdown.
    Returns a 'Drafting…' placeholder while no complete sections exist yet."""

    import re as _re

    FIELD_HEADERS = [
        ("title", None),
        ("problem_statement", "## Problem Statement"),
        ("target_users", "## Target Users"),
        ("user_stories", "## User Stories"),
        ("acceptance_criteria", "## Acceptance Criteria"),
        ("success_metrics", "## Success Metrics"),
        ("risks_open_questions", "## Risks & Open Questions"),
        ("out_of_scope", "## Out of Scope"),
    ]

    for closing in ['', '"}', '"]}', '}', ']}']:
        try:
            candidate = raw_json_partial.rstrip().rstrip(',') + closing
            parsed = json.loads(candidate)
            return _render_partial_prd(parsed, FIELD_HEADERS)
        except (json.JSONDecodeError, ValueError):
            continue

    rendered = []
    for field, header in FIELD_HEADERS:
        string_match = _re.search(
            rf'"{field}"\s*:\s*"((?:[^"\\]|\\.)*)"',
            raw_json_partial
        )
        if string_match:
            value = string_match.group(1).replace('\\n', '\n').replace('\\"', '"')
            if header:
                rendered.append(f"{header}\n\n{value}")
            else:
                rendered.append(f"# {value}")
            continue

        array_match = _re.search(
            rf'"{field}"\s*:\s*\[(.*?)\]',
            raw_json_partial,
            _re.DOTALL
        )
        if array_match:
            inner = array_match.group(1)
            items = _re.findall(r'"((?:[^"\\]|\\.)*)"', inner)
            if items and header:
                items_md = "\n".join(f"- {item}" for item in items)
                rendered.append(f"{header}\n\n{items_md}")

    if rendered:
        return "\n\n".join(rendered)
    return "*Drafting your PRD…*"


def _render_partial_prd(parsed: dict, field_headers: list) -> str:
    """Render whichever PRD fields are present in the parsed dict as markdown."""
    rendered = []
    for field, header in field_headers:
        if field not in parsed:
            continue
        value = parsed[field]
        if field == "title":
            rendered.append(f"# {value}")
        elif isinstance(value, str):
            rendered.append(f"{header}\n\n{value}")
        elif isinstance(value, list):
            if value and isinstance(value[0], dict):
                items_md = "\n".join(
                    f"- **{m.get('metric', '')}** — Baseline: {m.get('baseline', '')} | Target: {m.get('target', '')}"
                    for m in value if isinstance(m, dict)
                )
                rendered.append(f"{header}\n\n{items_md}")
            else:
                items_md = "\n".join(f"- {item}" for item in value)
                rendered.append(f"{header}\n\n{items_md}")
    return "\n\n".join(rendered) if rendered else "*Drafting your PRD…*"


def _prd_to_docx_bytes(prd: dict) -> bytes:
    doc = Document()
    doc.add_heading(prd.get("title", "Untitled"), level=1)

    doc.add_heading("Problem Statement", level=2)
    doc.add_paragraph(prd.get("problem_statement", ""))

    doc.add_heading("Target Users", level=2)
    for user in prd.get("target_users", []):
        doc.add_paragraph(user, style="List Bullet")

    doc.add_heading("User Stories", level=2)
    for story in prd.get("user_stories", []):
        doc.add_paragraph(story, style="List Number")

    doc.add_heading("Acceptance Criteria", level=2)
    for criterion in prd.get("acceptance_criteria", []):
        doc.add_paragraph(criterion, style="List Bullet")

    doc.add_heading("Success Metrics", level=2)
    metrics = prd.get("success_metrics", [])
    if metrics:
        table = doc.add_table(rows=len(metrics) + 1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Baseline", "Target"
        for i, m in enumerate(metrics, 1):
            row = table.rows[i].cells
            row[0].text = m.get("metric", "")
            row[1].text = m.get("baseline", "")
            row[2].text = m.get("target", "")

    doc.add_heading("Risks & Open Questions", level=2)
    for risk in prd.get("risks_open_questions", []):
        doc.add_paragraph(risk, style="List Bullet")

    doc.add_heading("Out of Scope", level=2)
    for item in prd.get("out_of_scope", []):
        doc.add_paragraph(item, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_full_prd(prd: dict) -> None:
    st.divider()
    st.header(prd.get("title", "Untitled"))

    st.subheader("Problem Statement")
    st.write(prd.get("problem_statement", ""))

    st.subheader("Target Users")
    for user in prd.get("target_users", []):
        st.markdown(f"- {user}")

    st.subheader("User Stories")
    for i, story in enumerate(prd.get("user_stories", []), 1):
        st.markdown(f"{i}. {story}")

    st.subheader("Acceptance Criteria")
    for criterion in prd.get("acceptance_criteria", []):
        st.markdown(f"- {criterion}")

    st.subheader("Success Metrics")
    metrics = prd.get("success_metrics", [])
    if metrics:
        cols = st.columns(3)
        cols[0].markdown("**Metric**")
        cols[1].markdown("**Baseline**")
        cols[2].markdown("**Target**")
        for m in metrics:
            c1, c2, c3 = st.columns(3)
            c1.write(m.get("metric", ""))
            c2.write(m.get("baseline", ""))
            c3.write(m.get("target", ""))

    st.subheader("Risks & Open Questions")
    for risk in prd.get("risks_open_questions", []):
        st.markdown(f"- {risk}")

    st.subheader("Out of Scope")
    for item in prd.get("out_of_scope", []):
        st.markdown(f"- {item}")

    _slug = _slugify(prd.get("title", "prd"))
    st.download_button(
        label="Download PRD",
        data=_prd_to_docx_bytes(prd),
        file_name=f"{_slug}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_{_slug}",
    )


def _render_eval(evaluation: dict) -> None:
    if not evaluation:
        return
    st.divider()
    st.header("Eval Results")
    overall = evaluation.get("overall_score", 0)
    color = "green" if overall >= 7.5 else "orange" if overall >= 5.0 else "red"
    st.markdown(
        f"<h2 style='color:{color}'>Overall Score: {overall} / 10</h2>",
        unsafe_allow_html=True,
    )
    st.caption(
        "Scores are directional quality signals, not absolute truth. "
        "The rubric is designed to highlight improvement areas and compare outputs consistently."
    )
    DIMENSIONS = ["specificity", "measurability", "testability", "user_centricity", "completeness"]
    scores = evaluation.get("scores", {})
    dim_cols = st.columns(5)
    for col, dim in zip(dim_cols, DIMENSIONS):
        entry = scores.get(dim, {})
        score = entry.get("score", "—")
        reason = entry.get("reason", "")
        col.metric(label=dim.replace("_", " ").title(), value=f"{score} / 10")
        col.caption(reason)
    st.subheader("Top 3 Improvements")
    for i, tip in enumerate(evaluation.get("top_3_improvements", []), 1):
        st.markdown(f"{i}. {tip}")


_SAMPLE_QUERIES = [
    "Create a PRD for a real-time ETA prediction and tracking experience for last-mile food delivery customers",
    "Create a PRD for a retail media insights dashboard for advertisers selling on a B2B e-commerce marketplace",
    "Create a PRD for a supplier-facing workflow tool that lets tour operators bulk-update inventory and pricing across marketplace listings",
]

st.set_page_config(page_title="PRD Generator", page_icon="📋", layout="wide")
st.title("PRD Generator with Eval Harness")
st.markdown("Generate structured PRDs from rough product ideas, then evaluate output quality with a calibrated LLM-as-judge rubric.")
st.caption(
    "_This is a portfolio demonstration of LLM evaluation methodology. "
    "The app uses paid API credits — please use thoughtfully. If you're a recruiter exploring my work, "
    "feel free to reach out via [LinkedIn](https://www.linkedin.com/in/sumit5suman/)._"
)

if "_pending_idea" in st.session_state:
    st.session_state["idea"] = st.session_state.pop("_pending_idea")

idea = st.text_area(
    "Describe your product idea",
    placeholder="e.g. a tool that helps remote engineering teams run async standups via Slack",
    height=120,
    key="idea",
)

st.markdown("**Try a sample idea:**")
_cc = st.columns(3)
for _i, (_col, _q) in enumerate(zip(_cc, _SAMPLE_QUERIES), 1):
    with _col:
        with st.container(border=True):
            st.markdown(_q)
            if st.button("Use this", key=f"chip_{_i}", use_container_width=True):
                st.session_state["_pending_idea"] = _q
                st.rerun()

if st.button("Generate PRD", type="primary", disabled=not idea.strip()):
    _retrieval_ok = True
    with st.spinner("Retrieving similar PRDs…"):
        try:
            retrieved_prds = get_similar_prd_metas(idea.strip(), k=2)
            retrieved_text = retrieve_similar_prds(idea.strip(), k=2)
        except Exception as e:
            st.error(f"Retrieval failed: {e}")
            _retrieval_ok = False

    if _retrieval_ok:
        for r in retrieved_prds:
            try:
                r["_why_retrieved"] = explain_retrieval(idea.strip(), r)
            except Exception:
                r["_why_retrieved"] = ""
        st.session_state["retrieved_prds"] = retrieved_prds
        st.session_state["_idea"] = idea.strip()
        st.session_state["_retrieved_examples"] = (
            "For reference, here are similar PRDs to use as style and structure "
            "guidance, not content to copy:\n\n" + retrieved_text
        )
        st.session_state.pop("prd", None)
        st.session_state.pop("evaluation", None)
        st.session_state["_generating"] = True
        st.rerun()

if "retrieved_prds" in st.session_state:
    with st.expander("Retrieved similar PRDs (RAG context)"):
        for r in st.session_state["retrieved_prds"]:
            st.markdown(f"**{r['title']}**")
            sentences = r.get("problem_statement", "").split(". ")
            summary = ". ".join(sentences[:2]) + ("." if len(sentences) > 2 else "")
            st.caption(summary)
            explanation = r.get("_why_retrieved", "")
            if explanation:
                st.caption(f"**Why retrieved:** {explanation}")

# Two placeholders: prd_placeholder owns PRD rendering, eval_placeholder owns eval scores.
prd_placeholder = st.empty()
eval_placeholder = st.empty()

# Streaming/eval phase — runs on the rerun immediately after button click.
# _generating is popped at the top so a mid-stream crash can't create an infinite loop.
if st.session_state.pop("_generating", False):
    _idea = st.session_state["_idea"]
    _examples = st.session_state["_retrieved_examples"]

    # Stage A: stream
    full_text = ""
    try:
        for chunk in generate_prd_streaming(_idea, _examples):
            full_text += chunk
            with prd_placeholder.container():
                st.markdown(_streaming_display_text(full_text))
    except Exception as e:
        prd_placeholder.error(f"PRD generation failed: {e}")
        st.stop()

    # Stage B: parse and render structured PRD once
    try:
        prd = parse_prd_text(full_text)
    except (json.JSONDecodeError, ValueError):
        prd_placeholder.error("Failed to parse generated PRD — output may have been truncated.")
        with st.expander("Raw output"):
            st.code(full_text)
        st.stop()

    st.session_state["prd"] = prd
    with prd_placeholder.container():
        _render_full_prd(prd)

    # Stage C: eval into its own placeholder
    with eval_placeholder.container():
        with st.spinner("Running LLM-as-judge evaluation…"):
            try:
                evaluation = evaluate_prd(prd)
            except Exception as e:
                st.error(f"Evaluation failed: {e}")
                evaluation = {}
        st.session_state["evaluation"] = evaluation
        _render_eval(evaluation)
        with st.expander("How it works (architecture)"):
            st.markdown(
                """
                1. **Idea input** — You paste a product idea in plain English.
                2. **Retrieval** — The system embeds your query using `sentence-transformers/all-MiniLM-L6-v2`
                   and retrieves the 2 most-similar PRDs from a curated corpus of 15 reference PRDs stored in Chroma.
                3. **Grounded generation** — Those retrieved PRDs are injected into the generation prompt
                   (Claude Sonnet 4.6), so the output inherits structural rigor without copying content.
                4. **LLM-as-judge eval** — The generated PRD is automatically scored 1–10 by a second
                   Claude call across 5 dimensions: specificity, measurability, testability, user-centricity,
                   and completeness. The rubric is calibrated to a strong/weak separation of 6.56 on a
                   20-PRD test set.

                Code, calibration details, and corpus on GitHub.
                """
            )
        st.caption(
            "Built by Sumit Suman · [LinkedIn](https://www.linkedin.com/in/sumit5suman/) · "
            "[GitHub](https://github.com/sumitsuman-droid/prd-generator-eval)"
        )

elif "prd" in st.session_state:
    # Subsequent reruns (e.g. Download button): re-render from session_state.
    with prd_placeholder.container():
        _render_full_prd(st.session_state["prd"])
    with eval_placeholder.container():
        _render_eval(st.session_state.get("evaluation", {}))
        with st.expander("How it works (architecture)"):
            st.markdown(
                """
                1. **Idea input** — You paste a product idea in plain English.
                2. **Retrieval** — The system embeds your query using `sentence-transformers/all-MiniLM-L6-v2`
                   and retrieves the 2 most-similar PRDs from a curated corpus of 15 reference PRDs stored in Chroma.
                3. **Grounded generation** — Those retrieved PRDs are injected into the generation prompt
                   (Claude Sonnet 4.6), so the output inherits structural rigor without copying content.
                4. **LLM-as-judge eval** — The generated PRD is automatically scored 1–10 by a second
                   Claude call across 5 dimensions: specificity, measurability, testability, user-centricity,
                   and completeness. The rubric is calibrated to a strong/weak separation of 6.56 on a
                   20-PRD test set.

                Code, calibration details, and corpus on GitHub.
                """
            )
        st.caption(
            "Built by Sumit Suman · [LinkedIn](https://www.linkedin.com/in/sumit5suman/) · "
            "[GitHub](https://github.com/sumitsuman-droid/prd-generator-eval)"
        )

if "prd" not in st.session_state and not st.session_state.get("_generating_active", False):
    with st.expander("How it works (architecture)"):
        st.markdown(
            """
            1. **Idea input** — You paste a product idea in plain English.
            2. **Retrieval** — The system embeds your query using `sentence-transformers/all-MiniLM-L6-v2`
               and retrieves the 2 most-similar PRDs from a curated corpus of 15 reference PRDs stored in Chroma.
            3. **Grounded generation** — Those retrieved PRDs are injected into the generation prompt
               (Claude Sonnet 4.6), so the output inherits structural rigor without copying content.
            4. **LLM-as-judge eval** — The generated PRD is automatically scored 1–10 by a second
               Claude call across 5 dimensions: specificity, measurability, testability, user-centricity,
               and completeness. The rubric is calibrated to a strong/weak separation of 6.56 on a
               20-PRD test set.

            Code, calibration details, and corpus on GitHub.
            """
        )
    st.caption(
        "Built by Sumit Suman · [LinkedIn](https://www.linkedin.com/in/sumit5suman/) · "
        "[GitHub](https://github.com/sumitsuman-droid/prd-generator-eval)"
    )
